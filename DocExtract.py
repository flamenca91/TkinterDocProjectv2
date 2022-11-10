from __future__ import annotations
import docx
from docx import Document
from docx.shared import RGBColor
import re
import xlwings

docFile = {"HRD":"HDS_new_pump.docx", "HRS":"HRS_new_pump.docx", "HTP":"HTP_new_pump.docx", "HTR":"HTR_new_pump.docx", \
           "PRS":"PRS_new_pump.docx", "RISK":"RiskAnalysis_Pump.docx", "SDS":"SDS_New_pump_x04.docx", \
           "ACE":"SRS_ACE_Pump_X01.docx", "BOLUS":"SRS_BolusCalc_Pump_X04.docx", "SRS":"SRS_DosingAlgorithm_X03.docx", \
           "SVAL":"SVaP_new_pump.docx", "SVATR":"SVaTR_new_pump.docx", "UT":"SVeTR_new_pump.docx", "URS":"URS_new_pump.docx"}

# DER and TBV are not valid tags
docRelation = {"HRD":("HRS"), "HRS":("PRS"), "PRS":("URS","RISK"), "HTR":("HTP"), "HTP":("HRD", "HRS"), \
               "SDS":("BOLUS","ACE","AID"), "ACE":("PRS"), "BOLUS":("PRS"), "AID":("PRS"), \
               "SVAL":("BOLUS", "ACE", "AID"), "SVATR":("SVAL"), "UT":("UNIT"), "INS": ("UNIT")}      # to be created by the GUI

filePath = "C:/Users/steph/OneDrive/Desktop/Docs_Project/"

docFileList = list(docFile.keys())                  # This is a list of all main tags found in each document
print(docFileList)
parentTagList = list(docRelation.values())          # List of all parent tags

report3 = Document()                #create word document
paragraph = report3.add_paragraph()
report3.save('report3.docx')

uniqueValidTagList = []                             # This is the valid child tag list
for tag in parentTagList:
    if type(tag) is tuple:                          # if a tuple is found, convert to a list and add to the list
        uniqueValidTagList.extend(list(tag))
    else:
        uniqueValidTagList.append(tag)              # if not a tuple simply append to the list
uniqueTagList = (list(set(uniqueValidTagList)))     # set() strips out all redundant tags

def GetText(filename):                      # Opens the document and places each paragraph into a list
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    fullText = [ele for ele in fullText if ele.strip()]   # Eliminates empty paragraphs
    return fullText

def GetParentTags():                    # Returns only valid parent tags
    for tag in docFileList:             # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            if tag == "BOLUS" or tag == "ACE":
                if re.search('.*[:\s]' + "SRS" + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\S*[:\s]' + "SRS" + '[:\s]\S*', t)
                    red = paragraph.add_run(y)
                    paragraph.add_run("\n\n")
                    red.bold = True
                    red.font.color.rgb = RGBColor(255, 0, 0)
                    #print(y[0])
                index = index + 1
            # print(ind)
            else:
                if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\S*[:\s]' + re.escape(tag) + '[:\s]\S*', t)
                    red = paragraph.add_run(y)
                    paragraph.add_run("\n\n")
                    red.bold = True
                    red.font.color.rgb = RGBColor(255, 0, 0)
                    #print(y[0])
                index = index + 1
            #print(ind)

def GetChildTags():                     # Returns only valid child tags
    for tag in docFileList:             # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        unique_tags = []
        index = 0
        ind = []
        for t in textList:
            if tag == "BOLUS" or tag == "ACE":
                if re.search('.*[:\s]' + "SRS" + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('[\[{].+[\]}]', t)
                    if len(y) != 0:
                        unique_tags.append(y[0])
                        green = paragraph.add_run(y[0])
                        paragraph.add_run("\n\n")
                        green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                        green.bold = True
                        #print(y[0])
                index = index + 1
                # print(ind)
            else:
                if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('[\[{].+[\]}]', t)
                    if len(y) != 0:
                        unique_tags.append(y[0])
                        green = paragraph.add_run(y[0])
                        paragraph.add_run("\n\n")
                        green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                        green.bold = True
                        #print(y[0])
                index = index + 1
            #print(ind)


def GetOrphanTags():
    for tag in docFileList:  # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            #y = re.findall('[\s\]]\[.+\][\[\s]', t)
            y = re.findall('\[.+\]', t)
            if len(y) != 0:

                #green = paragraph.add_run(y[0])
                #paragraph.add_run("\n\n")
                #green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                #green.bold = True

                print(y[0])
                index = index + 1
                # print(ind)

            #else:
            #    if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
            #        ind.append(index)
            #        tt = t
            #        y = re.findall('\s\[.+\]\s', t)
            #        if len(y) != 0:
            #            green = paragraph.add_run(y[0])
            #            paragraph.add_run("\n\n")
            #            green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            #            green.bold = True
            #            # print(y[0])
            #    index = index + 1
            # print(ind)

'''
runner2 = paragraph.add_run("\n\nParent tag/tags\n\n")
runner2.bold = True                              #make it bold
GetParentTags()

runner2 = paragraph.add_run("\n\nChild tag/tags\n\n")
runner2.bold = True
GetChildTags()

#runner2 = paragraph.add_run("\n\nOrphanChild tag/tags\n")
#runner2.bold = True
#GetOrphanTags()

report3.save('report3.docx')
GetOrphanTags()
'''